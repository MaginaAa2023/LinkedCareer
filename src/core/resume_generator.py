from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

class ResumeGenerator:
    def __init__(self):
        # 支持的模板列表
        self.templates = {
            'executive': '旗舰商务版（适合高管/金融/国企）',
            'modern': '现代双栏版（适合互联网/科技行业）',
            'soe': '国企稳重版（适合央企/事业单位/体制内）'
        }
    
    def set_font(self, run, size=10.5, bold=False, color=None):
        """设置全局字体"""
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    
    def set_cell_bg(self, cell, rgb_hex):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), rgb_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def hide_cell_border(self, cell):
        """隐藏单元格边框"""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            borderElm = OxmlElement(f'w:{border}')
            borderElm.set(qn('w:val'), 'nil')
            tcBorders.append(borderElm)
        tcPr.append(tcBorders)
    
    def generate_executive(self, data, output_path):
        """生成旗舰商务版简历"""
        doc = Document()
        sections = doc.sections[0]
        sections.top_margin = Cm(2)
        sections.bottom_margin = Cm(2)
        sections.left_margin = Cm(2)
        sections.right_margin = Cm(2)

        # 头部通栏背景
        table = doc.add_table(rows=1, cols=2)
        table.width = Cm(17)
        table.autofit = False
        table.columns[0].width = Cm(13)
        table.columns[1].width = Cm(4)
        for row in table.rows:
            for cell in row.cells:
                self.hide_cell_border(cell)
                self.set_cell_bg(cell, '003866')  # 深蓝背景
        
        # 左侧文字
        left_cell = table.cell(0, 0)
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = left_cell.paragraphs[0].add_run(data['basicInfo']['name'])
        self.set_font(run, size=18, bold=True, color=(255,255,255))

        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['basicInfo']['position'])
        self.set_font(run, size=12, color=(255,255,255))

        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = f"{data['basicInfo']['phone']} | {data['basicInfo']['email']} | {data['basicInfo']['city']} | {data['basicInfo']['degree']}"
        run = p.add_run(contact)
        self.set_font(run, size=10, color=(255,255,255))

        # 右侧头像
        right_cell = table.cell(0, 1)
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = right_cell.paragraphs[0].add_run("□□□□□\n□ 头像 □\n□□□□□")
        self.set_font(run, size=9, color=(200,200,200))

        doc.add_paragraph()
        p = doc.add_paragraph("—" * 70)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(p.runs[0], size=9, color=(0, 86, 145))

        # 核心优势
        p = doc.add_paragraph()
        run = p.add_run("◆ 核心优势")
        self.set_font(run, size=12, bold=True, color=(0, 56, 102))
        p = doc.add_paragraph("—" * 30)
        self.set_font(p.runs[0], size=9, color=(0, 86, 145))

        for advantage in data['advantages']:
            p = doc.add_paragraph(f"• {advantage}")
            self.set_font(p.runs[0], size=10.5)
        
        doc.add_paragraph()

        # 工作经历
        p = doc.add_paragraph()
        run = p.add_run("■ 工作经历")
        self.set_font(run, size=12, bold=True, color=(0, 56, 102))
        p = doc.add_paragraph("—" * 30)
        self.set_font(p.runs[0], size=9, color=(0, 86, 145))

        for exp in data['experiences']:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp['company']} · {exp['position']}")
            self.set_font(run, size=11, bold=True)
            run = p.add_run(f"{' ' * (40 - len(exp['company'] + exp['position']))}{exp['time']}")
            self.set_font(run, size=10, color=(80,80,80))

            p = doc.add_paragraph(exp['description'])
            self.set_font(p.runs[0], size=10.5)

            for ach in exp['achievements']:
                p = doc.add_paragraph(f"• {ach}")
                self.set_font(p.runs[0], size=10.5)
            doc.add_paragraph()
        
        # 核心技能
        p = doc.add_paragraph()
        run = p.add_run("● 核心技能")
        self.set_font(run, size=12, bold=True, color=(0, 56, 102))
        p = doc.add_paragraph("—" * 30)
        self.set_font(p.runs[0], size=9, color=(0, 86, 145))

        for category, skills in data['skills'].items():
            p = doc.add_paragraph()
            run = p.add_run(f"• {category}：")
            self.set_font(run, size=10.5, bold=True)
            run = p.add_run("、".join(skills))
            self.set_font(run, size=10.5)
        
        doc.add_paragraph()

        # 教育背景
        p = doc.add_paragraph()
        run = p.add_run("▲ 教育背景")
        self.set_font(run, size=12, bold=True, color=(0, 56, 102))
        p = doc.add_paragraph("—" * 30)
        self.set_font(p.runs[0], size=9, color=(0, 86, 145))

        for edu in data['education']:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu['school']} · {edu['major']} · {edu['degree']}")
            self.set_font(run, size=11, bold=True)
            run = p.add_run(f"{' ' * (40 - len(edu['school'] + edu['major'] + edu['degree']))}{edu['time']}")
            self.set_font(run, size=10, color=(80,80,80))

        doc.save(output_path)
        return output_path
    
    def generate_modern(self, data, output_path):
        """生成现代双栏版简历"""
        doc = Document()
        sections = doc.sections[0]
        sections.top_margin = Cm(1.5)
        sections.bottom_margin = Cm(1.5)
        sections.left_margin = Cm(0)
        sections.right_margin = Cm(0)

        # 整体两栏布局
        main_table = doc.add_table(rows=1, cols=2)
        main_table.width = Cm(21)
        main_table.autofit = False
        main_table.columns[0].width = Cm(6.3)
        main_table.columns[1].width = Cm(14.7)

        left_cell = main_table.cell(0, 0)
        right_cell = main_table.cell(0, 1)
        self.hide_cell_border(left_cell)
        self.hide_cell_border(right_cell)
        self.set_cell_bg(left_cell, 'F0F2F5')

        # 左侧侧边栏
        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("■■■■■\n■ 头像 ■\n■■■■■")
        self.set_font(run, size=12, color=(150,150,150))

        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['basicInfo']['name'])
        self.set_font(run, size=16, bold=True, color=(30, 136, 229))

        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['basicInfo']['position'])
        self.set_font(run, size=11, color=(80,80,80))

        p = left_cell.add_paragraph("—"*20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(p.runs[0], size=8, color=(180,180,180))

        # 基本信息
        p = left_cell.add_paragraph()
        contact = f"📱 {data['basicInfo']['phone']}\n📧 {data['basicInfo']['email']}\n📍 {data['basicInfo']['city']}\n🎓 {data['basicInfo']['degree']}\n"
        run = p.add_run(contact)
        self.set_font(run, size=9.5)

        # 核心标签
        p = left_cell.add_paragraph()
        run = p.add_run("核心标签")
        self.set_font(run, size=11, bold=True, color=(30, 136, 229))
        tags_str = "\n".join([f"• {tag}" for tag in data['tags']])
        p = left_cell.add_paragraph(tags_str)
        self.set_font(p.runs[0], size=9.5)

        # 核心技能
        p = left_cell.add_paragraph()
        run = p.add_run("核心技能")
        self.set_font(run, size=11, bold=True, color=(30, 136, 229))
        for category, skills in data['skills'].items():
            p = left_cell.add_paragraph()
            run = p.add_run(category)
            self.set_font(run, size=9.5, bold=True)
            p = left_cell.add_paragraph("、".join(skills[:4]) + "\n")
            self.set_font(p.runs[0], size=9)

        # 教育背景
        p = left_cell.add_paragraph()
        run = p.add_run("教育背景")
        self.set_font(run, size=11, bold=True, color=(30, 136, 229))
        edu = data['education'][0]
        edu_str = f"{edu['school']}\n{edu['major']} {edu['degree']}\n{edu['time']}"
        p = left_cell.add_paragraph(edu_str)
        self.set_font(p.runs[0], size=9.5)

        # 右侧内容区
        # 核心优势
        p = right_cell.paragraphs[0]
        run = p.add_run("核心优势")
        self.set_font(run, size=14, bold=True, color=(30, 136, 229))
        p = right_cell.add_paragraph("—"*50)
        self.set_font(p.runs[0], size=9, color=(30, 136, 229))

        advantages_str = "\n".join([f"• {adv}" for adv in data['advantages']]) + "\n"
        p = right_cell.add_paragraph(advantages_str)
        self.set_font(p.runs[0], size=10.5)

        # 工作经历
        p = right_cell.add_paragraph()
        run = p.add_run("工作经历")
        self.set_font(run, size=14, bold=True, color=(30, 136, 229))
        p = right_cell.add_paragraph("—"*50)
        self.set_font(p.runs[0], size=9, color=(30, 136, 229))

        for exp in data['experiences']:
            p = right_cell.add_paragraph()
            run = p.add_run(f"{exp['company']} · {exp['position']}")
            self.set_font(run, size=12, bold=True)
            run = p.add_run(f"{' ' * (35 - len(exp['company'] + exp['position']))}{exp['time']}")
            self.set_font(run, size=10, color=(80,80,80))

            p = right_cell.add_paragraph(exp['description'])
            self.set_font(p.runs[0], size=10)

            ach_str = "\n".join([f"• {ach}" for ach in exp['achievements']]) + "\n"
            p = right_cell.add_paragraph(ach_str)
            self.set_font(p.runs[0], size=10)

        doc.save(output_path)
        return output_path
    
    def generate_soe(self, data, output_path):
        """生成国企稳重版简历，适合央企/事业单位/体制内"""
        doc = Document()
        sections = doc.sections[0]
        sections.top_margin = Cm(2.5)
        sections.bottom_margin = Cm(2.5)
        sections.left_margin = Cm(2.5)
        sections.right_margin = Cm(2.5)

        # 头部：姓名+求职意向，黑色加粗，正式
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['basicInfo']['name'])
        self.set_font(run, size=16, bold=True, color=(0,0,0))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['basicInfo']['position'])
        self.set_font(run, size=12, color=(51,51,51))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = f"电话：{data['basicInfo']['phone']} | 邮箱：{data['basicInfo']['email']} | 籍贯：{data['basicInfo'].get('native_place','')} | 政治面貌：{data['basicInfo'].get('political_status','群众')}"
        run = p.add_run(contact)
        self.set_font(run, size=10.5, color=(51,51,51))

        p = doc.add_paragraph("_" * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(p.runs[0], size=9, color=(102,102,102))
        doc.add_paragraph()

        # 教育背景（国企优先看教育，放在最前面）
        p = doc.add_paragraph()
        run = p.add_run("教育背景")
        self.set_font(run, size=12, bold=True, color=(0,0,0))
        p = doc.add_paragraph("—" * 25)
        self.set_font(p.runs[0], size=9, color=(0,0,0))

        for edu in data['education']:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu['school']}   {edu['major']}   {edu['degree']}")
            self.set_font(run, size=11, bold=True)
            run = p.add_run(f"{' ' * 30}{edu['time']}")
            self.set_font(run, size=10, color=(51,51,51))
            doc.add_paragraph()

        # 工作经历
        p = doc.add_paragraph()
        run = p.add_run("工作经历")
        self.set_font(run, size=12, bold=True, color=(0,0,0))
        p = doc.add_paragraph("—" * 25)
        self.set_font(p.runs[0], size=9, color=(0,0,0))

        for exp in data['experiences']:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp['company']}   {exp['position']}")
            self.set_font(run, size=11, bold=True)
            run = p.add_run(f"{' ' * 30}{exp['time']}")
            self.set_font(run, size=10, color=(51,51,51))

            p = doc.add_paragraph(exp['description'])
            self.set_font(p.runs[0], size=10.5)

            for ach in exp['achievements']:
                p = doc.add_paragraph(f"• {ach}")
                self.set_font(p.runs[0], size=10.5)
            doc.add_paragraph()

        # 专业技能/证书
        p = doc.add_paragraph()
        run = p.add_run("专业技能与证书")
        self.set_font(run, size=12, bold=True, color=(0,0,0))
        p = doc.add_paragraph("—" * 25)
        self.set_font(p.runs[0], size=9, color=(0,0,0))

        for category, skills in data['skills'].items():
            p = doc.add_paragraph()
            run = p.add_run(f"• {category}：")
            self.set_font(run, size=10.5, bold=True)
            run = p.add_run("、".join(skills))
            self.set_font(run, size=10.5)
        doc.add_paragraph()

        # 获奖情况/资质
        p = doc.add_paragraph()
        run = p.add_run("获奖情况与资质")
        self.set_font(run, size=12, bold=True, color=(0,0,0))
        p = doc.add_paragraph("—" * 25)
        self.set_font(p.runs[0], size=9, color=(0,0,0))

        for honor in data.get('honors', []):
            p = doc.add_paragraph(f"• {honor}")
            self.set_font(p.runs[0], size=10.5)

        doc.save(output_path)
        return output_path
    
    def generate_pdf(self, data, template='executive', output_path='resume.pdf'):
        """生成PDF格式简历，和Word样式1:1一致"""
        # 先生成临时Word文件
        temp_docx = output_path.replace('.pdf', '_temp.docx')
        self.generate(data, template, temp_docx)
        
        # 转换为PDF
        from docx2html import convert
        import weasyprint
        
        # Word转HTML
        html_content = convert(temp_docx)
        
        # 添加打印样式，保证和Word一致
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                    @top-center {{
                        content: "{data['basicInfo']['name']} - 简历;
                        font-size: 10pt;
                        color: #666;
                    }}
                    @bottom-center {{
                        content: "第 " counter(page) " 页 / 共 " counter(pages) " 页";
                        font-size: 9pt;
                        color: #666;
                    }}
                }}
                body {{
                    font-family: "微软雅黑", "Microsoft YaHei", sans-serif;
                    font-size: 10.5pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1, h2, h3 {{
                    font-weight: bold;
                }}
                .highlight {{
                    color: #c84242;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 生成PDF
        weasyprint.HTML(string=styled_html).write_pdf(output_path)
        
        # 清理临时文件
        import os
        os.remove(temp_docx)
        
        return output_path
    
    def generate(self, data, template='executive', output_path='resume.docx', format='docx'):
        """生成简历入口"""
        if format == 'pdf':
            return self.generate_pdf(data, template, output_path)
        elif format == 'docx':
            if template == 'executive':
                return self.generate_executive(data, output_path)
            elif template == 'modern':
                return self.generate_modern(data, output_path)
            elif template == 'soe':
                return self.generate_soe(data, output_path)
            else:
                raise ValueError(f"不支持的模板类型：{template}，支持的模板：{list(self.templates.keys())}")
        else:
            raise ValueError(f"不支持的输出格式：{format}，支持的格式：docx、pdf")
