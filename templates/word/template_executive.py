from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 1. 创建文档，设置基础样式
doc = Document()
sections = doc.sections[0]
sections.top_margin = Cm(2)
sections.bottom_margin = Cm(2)
sections.left_margin = Cm(2)
sections.right_margin = Cm(2)

# 设置全局字体
def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# 设置段落背景色
def set_paragraph_background(p, rgb_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), rgb_color)
    p._p.get_or_add_pPr().append(shading_elm)

# ----------------------
# 头部：通栏彩色背景+反白字体+头像悬浮（两栏布局）
# ----------------------
# 创建无框两列表格，通栏宽度
table = doc.add_table(rows=1, cols=2)
table.width = Cm(17)  # 整个表格宽度和页面一致
table.autofit = False
table.columns[0].width = Cm(13)  # 左侧文字区域
table.columns[1].width = Cm(4)   # 右侧头像区域
# 隐藏表格边框
for row in table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            borderElm = OxmlElement(f'w:{border}')
            borderElm.set(qn('w:val'), 'nil')
            tcBorders.append(borderElm)
        tcPr.append(tcBorders)
        # 所有单元格统一深蓝背景，实现通栏效果
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003866')  # 统一深蓝背景
        cell._tc.get_or_add_tcPr().append(shading)

# 左侧：姓名+职位+联系方式，白字
left_cell = table.cell(0, 0)
left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = left_cell.paragraphs[0].add_run("郝霄")
set_font(run, size=18, bold=True, color=(255, 255, 255))  # 白色字体

p = left_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("企业数字化负责人 / 高级产品总监 / 资管产品负责人")
set_font(run, size=12, color=(255, 255, 255))

p = left_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("186 1162 6420 | airaurum@gmail.com | 北京市 | 本科")
set_font(run, size=10, color=(255, 255, 255))

# 右侧：头像占位，悬浮在蓝色背景上
right_cell = table.cell(0, 1)
right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# 灰色头像占位框
run = right_cell.paragraphs[0].add_run("□□□□□\n□ 头像 □\n□□□□□")
set_font(run, size=9, color=(200, 200, 200))
p = right_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2寸商务照")
set_font(run, size=8, color=(200, 200, 200))

doc.add_paragraph()
# 分割线
p = doc.add_paragraph("—" * 70)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.runs[0], size=9, color=(0, 86, 145))

# ----------------------
# 核心优势模块
# ----------------------
p = doc.add_paragraph()
run = p.add_run("◆ 核心优势")
set_font(run, size=12, bold=True, color=(0, 56, 102))

p = doc.add_paragraph()
run = p.add_run("• 2020年入选龙湖集团正式合伙人，2024-2025连续两年获评高级合伙人，位列集团2万员工前80名\n")
set_font(run, size=10.5)
run = p.add_run("• 带领150人团队，支撑龙湖全业态年交易流水超百亿，智慧物业数字化建设年降本超10亿元\n")
set_font(run, size=10.5)
run = p.add_run("• 主导百度内容营销产品线，带领10人团队实现年度收入突破10亿大关\n")
set_font(run, size=10.5)
run = p.add_run("• 具备0到1产品搭建、百人团队管理、数字化战略规划、商业化拓展全链路能力")
set_font(run, size=10.5)

doc.add_paragraph()

# ----------------------
# 工作经历模块
# ----------------------
p = doc.add_paragraph()
run = p.add_run("■ 工作经历")
set_font(run, size=12, bold=True, color=(0, 56, 102))
p = doc.add_paragraph("—" * 30)
set_font(p.runs[0], size=9, color=(0, 86, 145))

# 千丁数科
p = doc.add_paragraph()
run = p.add_run("千丁数科 · 智慧资管部负责人")
set_font(run, size=11, bold=True)
run = p.add_run("                          2026.01 - 至今")
set_font(run, size=10, color=(80, 80, 80))

p = doc.add_paragraph()
run = p.add_run("龙湖数字科技部资管相关业务线独立为千丁数科智慧资管部，作为部门负责人负责对内服务龙湖原有业务，对外开展商业化资管系统和AI Agent产品服务：")
set_font(run, size=10.5)

p = doc.add_paragraph("• 负责龙湖天街、冠寓核心业务系统建设与迭代\n• 主导不动产资产管理与运营相关系统和AI Agent产品研发\n• 面向国央企客户提供私有化部署、定制开发、实施运维全链路服务")
set_font(p.runs[0], size=10.5)

doc.add_paragraph()

# 龙湖集团
p = doc.add_paragraph()
run = p.add_run("龙湖集团数字科技部 · 中心负责人")
set_font(run, size=11, bold=True)
run = p.add_run("                     2019.11 - 2025.12")
set_font(run, size=10, color=(80, 80, 80))

p = doc.add_paragraph()
run = p.add_run("先后负责办公协同、智慧物业数字化、全业态C端产品等核心产品线，带领团队规模30-150人，支撑年交易流水超百亿：")
set_font(run, size=10.5)

p = doc.add_paragraph()
run = p.add_run("• 2020年入选龙湖合伙人，2024-2025获评高级合伙人，集团2万员工前80\n")
set_font(run, size=10.5)
run = p.add_run("• 2023.10 - 2025.12 数科智能引擎客户应用中心/智慧服务中心负责人：服务会员超9000万，年交易流水超百亿\n")
set_font(run, size=10.5)
run = p.add_run("• 2021.06 - 2023.10 智慧物业数字化中心负责人：支撑年营收120亿，年降本超")
set_font(run, size=10.5)
run = p.add_run(" 10亿元")
set_font(run, size=10.5, bold=True, color=(194, 62, 62))
run = p.add_run("\n• 2019.11 - 2021.06 办公协同中心产品线负责人：主导自研云图梭协同办公平台，支撑3.5万员工+10万供方使用\n• 创立数科产品委员会并任组长，完善百人产品团队能力体系\n• 全业态会员体系整合：年积分流通量超5亿元，带动消费场景数亿元营收")
set_font(run, size=10.5)

doc.add_paragraph()

# 百度
p = doc.add_paragraph()
run = p.add_run("百度 · 搜索商业化高级产品经理")
set_font(run, size=11, bold=True)
run = p.add_run("                       2012.08 - 2019.11")
set_font(run, size=10, color=(80, 80, 80))

p = doc.add_paragraph()
run = p.add_run("负责百度内容营销产品线全流程产品工作，从产品运营专员成长为高级产品经理，带领10人团队实现年度收入突破10亿大关：")
set_font(run, size=10.5)

p = doc.add_paragraph()
run = p.add_run("• 内容营销整合破冰项目（2019）：日均收入从25万提升至")
set_font(run, size=10.5)
run = p.add_run(" 150万")
set_font(run, size=10.5, bold=True, color=(194, 62, 62))
run = p.add_run("\n• 商业阿拉丁售卖机制革新项目（2017）：3C、汽车垂类收入增长超")
set_font(run, size=10.5)
run = p.add_run(" 60%")
set_font(run, size=10.5, bold=True, color=(194, 62, 62))
run = p.add_run("\n• 品牌专区高级感/动态物料项目（2016）：平均点击率提升")
set_font(run, size=10.5)
run = p.add_run(" 25%")
set_font(run, size=10.5, bold=True, color=(194, 62, 62))
run = p.add_run("，1年内80%客户选择升级，覆盖年收入35亿")
set_font(run, size=10.5)

doc.add_paragraph()

# 国双科技
p = doc.add_paragraph()
run = p.add_run("国双科技 · 数字化营销小组主管")
set_font(run, size=11, bold=True)
run = p.add_run("                       2011.06 - 2012.08")
set_font(run, size=10, color=(80, 80, 80))

p = doc.add_paragraph()
run = p.add_run("服务微软云、Agoda、Opera、开心网等客户，负责搜索引擎广告投放：")
set_font(run, size=10.5)

p = doc.add_paragraph("• 直接负责年度预算超8000万元，管理关键词百万量级，6个月快速晋升为小组主管\n• 参与设计国双SEM Dissector系统，对接百度推广API，实现MVP上线和持续迭代")
set_font(p.runs[0], size=10.5)

doc.add_paragraph()

# ----------------------
# 核心技能模块
# ----------------------
p = doc.add_paragraph()
run = p.add_run("● 核心技能")
set_font(run, size=12, bold=True, color=(0, 56, 102))
p = doc.add_paragraph("—" * 30)
set_font(p.runs[0], size=9, color=(0, 86, 145))

p = doc.add_paragraph()
run = p.add_run("• 产品能力：")
set_font(run, size=10.5, bold=True)
run = p.add_run("产品战略规划 · 0到1产品搭建 · 团队管理 · 跨部门协作 · 数字化转型 · 商业化产品 · 用户增长 · 收入提升 · 降本提效")
set_font(run, size=10.5)

p = doc.add_paragraph()
run = p.add_run("• 行业经验：")
set_font(run, size=10.5, bold=True)
run = p.add_run("搜索商业化 · 内容营销 · 智慧物业 · 协同办公 · 会员体系 · 企业IT建设 · 消费互联网 · 产业互联网 · 不动产资管")
set_font(run, size=10.5)

p = doc.add_paragraph()
run = p.add_run("• 技术理解：")
set_font(run, size=10.5, bold=True)
run = p.add_run("SaaS产品 · 移动端APP · 小程序 · 系统架构 · 数据运营 · IoT智能设备 · IM即时通讯 · 音视频会议 · AI Agent")
set_font(run, size=10.5)

p = doc.add_paragraph()
run = p.add_run("• 管理能力：")
set_font(run, size=10.5, bold=True)
run = p.add_run("百人团队管理 · 预算管理 · 项目管理 · 敏捷开发 · 产品体系建设")
set_font(run, size=10.5)

doc.add_paragraph()

# ----------------------
# 教育背景模块
# ----------------------
p = doc.add_paragraph()
run = p.add_run("▲ 教育背景")
set_font(run, size=12, bold=True, color=(0, 56, 102))
p = doc.add_paragraph("—" * 30)
set_font(p.runs[0], size=9, color=(0, 86, 145))

p = doc.add_paragraph()
run = p.add_run("天津科技大学 · 通信工程 · 本科")
set_font(run, size=11, bold=True)
run = p.add_run("                          2007.09 - 2011.06")
set_font(run, size=10, color=(80, 80, 80))

# 保存文档
output_path = "郝霄_高管商务简历_最终版.docx"
doc.save(output_path)
print(f"✅ 生成成功：{output_path}")
